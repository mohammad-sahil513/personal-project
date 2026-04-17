"""
Deterministic DOCX structure extractor for custom template compilation.

Phase 6 scope:
- extract heading-like paragraphs from a DOCX file,
- capture basic document signals required by later compiler/layout phases,
- avoid any AI/LLM logic.

Extraction policy:
- prefer explicit Word heading styles,
- allow a bounded numbered-heading fallback for reasonably structured documents,
- capture only lightweight structural metadata here.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from ..contracts.compiler_contracts import ExtractedDocxStructure, ExtractedHeading
from .header_normalizer import HeaderNormalizer


class DocxExtractor:
    """Extract deterministic heading structure and coarse document signals from DOCX."""

    _HEADING_STYLE_PATTERN = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)
    _NUMBERED_HEADING_PATTERN = re.compile(r"^\s*(\d+(?:\.\d+)*)[\)\.\-]?\s+(.+?)\s*$")

    def __init__(self, *, header_normalizer: HeaderNormalizer | None = None) -> None:
        self._header_normalizer = header_normalizer or HeaderNormalizer()

    def extract(self, docx_path: str | Path) -> ExtractedDocxStructure:
        """
        Extract headings and coarse structure flags from a DOCX template.

        Args:
            docx_path: Path to the DOCX template.

        Returns:
            ExtractedDocxStructure with deterministic heading metadata.
        """
        path = Path(docx_path)
        document = Document(path)

        headings: list[ExtractedHeading] = []
        heading_order = 0

        for paragraph in document.paragraphs:
            raw_text = paragraph.text.strip()
            if not raw_text:
                continue

            heading_level = self._infer_heading_level(paragraph)
            if heading_level is None:
                continue

            normalized_text = self._header_normalizer.normalize(raw_text)
            if not normalized_text:
                continue

            headings.append(
                ExtractedHeading(
                    raw_text=raw_text,
                    normalized_text=normalized_text,
                    level=heading_level,
                    order_index=heading_order,
                )
            )
            heading_order += 1

        return ExtractedDocxStructure(
            headings=headings,
            contains_tables=bool(document.tables),
            contains_headers_footers=self._contains_headers_or_footers(document),
            contains_multiple_sections=len(document.sections) > 1,
        )

    def _infer_heading_level(self, paragraph: Paragraph) -> int | None:
        """
        Infer heading level from paragraph metadata.

        Strategy:
        1. explicit Word heading style,
        2. bounded numbered-heading fallback.
        """
        style_name = (paragraph.style.name or "").strip() if paragraph.style is not None else ""
        style_match = self._HEADING_STYLE_PATTERN.match(style_name)
        if style_match:
            return max(1, min(9, int(style_match.group(1))))

        text = paragraph.text.strip()
        numbered_match = self._NUMBERED_HEADING_PATTERN.match(text)
        if numbered_match:
            if self._looks_like_reasonable_fallback_heading(text):
                numbering = numbered_match.group(1)
                return max(1, min(9, numbering.count(".") + 1))

        return None

    @staticmethod
    def _looks_like_reasonable_fallback_heading(text: str) -> bool:
        """
        Guard the numbered-heading fallback so normal paragraphs are not misclassified.
        """
        words = text.split()
        return len(words) <= 15 and len(text) <= 140

    @staticmethod
    def _contains_headers_or_footers(document: DocxDocument) -> bool:
        """Return True if any section contains non-empty header/footer text."""
        for section in document.sections:
            header_text = "\n".join(
                para.text.strip()
                for para in section.header.paragraphs
                if para.text and para.text.strip()
            )
            footer_text = "\n".join(
                para.text.strip()
                for para in section.footer.paragraphs
                if para.text and para.text.strip()
            )
            if header_text or footer_text:
                return True

        return False