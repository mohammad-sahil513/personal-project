"""
Shell builder for custom DOCX templates.

This builder creates a cleared shell DOCX that preserves:
- section breaks,
- page setup,
- headers/footers,
- overall document structure,

while removing main-body content so later rendering can inject generated
content into a corporate template shell.
"""

from __future__ import annotations

import logging
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document


@dataclass(frozen=True, slots=True)
class ShellBuildResult:
    """Summary of one shell-build operation."""

    shell_docx_path: Path
    cleared_paragraph_count: int
    cleared_table_cell_count: int


class ShellBuilder:
    """Create a cleared shell DOCX from a source custom template."""

    def __init__(
        self,
        *,
        output_root: str | Path | None = None,
        logger: logging.Logger | None = None,
    ) -> None:
        resolved_output_root = (
            Path(output_root).resolve()
            if output_root is not None
            else Path(__file__).resolve().parents[4] / "artifacts" / "template" / "shells"
        )
        self._output_root = resolved_output_root
        self._logger = logger or logging.getLogger(__name__)

    @property
    def output_root(self) -> Path:
        """Return the resolved shell output root directory."""
        return self._output_root

    def build_shell(
        self,
        *,
        source_docx_path: str | Path,
        template_id: str,
        version: str,
        output_filename: str | None = None,
    ) -> ShellBuildResult:
        """
        Build a cleared shell DOCX from the source custom template.
        """
        source_path = Path(source_docx_path).resolve()
        output_dir = self._output_root / template_id / version
        output_dir.mkdir(parents=True, exist_ok=True)

        filename = output_filename or f"{template_id}_{version}_shell.docx"
        shell_path = (output_dir / filename).resolve()

        self._log_info(
            "shell_build_start",
            template_id=template_id,
            template_version=version,
            source_docx_path=str(source_path),
            shell_docx_path=str(shell_path),
        )

        shutil.copy2(source_path, shell_path)
        document = Document(shell_path)

        cleared_paragraph_count = 0
        for paragraph in document.paragraphs:
            if paragraph.text:
                self._clear_paragraph(paragraph)
                cleared_paragraph_count += 1

        cleared_table_cell_count = 0
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        for paragraph in cell.paragraphs:
                            self._clear_paragraph(paragraph)
                        cleared_table_cell_count += 1

        document.save(shell_path)

        self._log_info(
            "shell_build_completed",
            template_id=template_id,
            template_version=version,
            shell_docx_path=str(shell_path),
            cleared_paragraph_count=cleared_paragraph_count,
            cleared_table_cell_count=cleared_table_cell_count,
        )

        return ShellBuildResult(
            shell_docx_path=shell_path,
            cleared_paragraph_count=cleared_paragraph_count,
            cleared_table_cell_count=cleared_table_cell_count,
        )

    @staticmethod
    def _clear_paragraph(paragraph) -> None:
        """
        Clear paragraph text while preserving the paragraph container and style.
        """
        for run in paragraph.runs:
            run.text = ""

    def _log_info(self, event_name: str, **payload: object) -> None:
        """Emit a lightweight structured-ish log entry."""
        self._logger.info("%s | %s", event_name, payload)