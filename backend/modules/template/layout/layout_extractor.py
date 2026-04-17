"""
Layout extractor for custom DOCX templates.

This is the main Template-side entry point for building a versioned
LayoutManifest from a corporate DOCX template.

Responsibilities:
- load DOCX,
- parse page setup, headers/footers, styles, and tables,
- extract heading-like anchors,
- build triple-anchor-friendly metadata,
- return a strongly typed LayoutManifest.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from ..compiler.header_normalizer import HeaderNormalizer
from .header_footer_parser import HeaderFooterParser
from .layout_contracts import AnchorMetadata, LayoutManifest
from .page_setup_parser import PageSetupParser
from .style_parser import StyleParser
from .table_format_parser import TableFormatParser


class LayoutExtractor:
    """Extract a versioned LayoutManifest from a custom DOCX template."""

    _HEADING_STYLE_PATTERN = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)
    _NUMBERED_HEADING_PATTERN = re.compile(r"^\s*(\d+(?:\.\d+)*)[\)\.\-]?\s+(.+?)\s*$")
    _PARA_ID_ATTR = "{http://schemas.microsoft.com/office/word/2010/wordml}paraId"

    def __init__(
        self,
        *,
        page_setup_parser: PageSetupParser | None = None,
        header_footer_parser: HeaderFooterParser | None = None,
        style_parser: StyleParser | None = None,
        table_format_parser: TableFormatParser | None = None,
        header_normalizer: HeaderNormalizer | None = None,
        logger: logging.Logger | None = None,
    ) -> None:
        self._page_setup_parser = page_setup_parser or PageSetupParser()
        self._header_footer_parser = header_footer_parser or HeaderFooterParser()
        self._style_parser = style_parser or StyleParser()
        self._table_format_parser = table_format_parser or TableFormatParser()
        self._header_normalizer = header_normalizer or HeaderNormalizer()
        self._logger = logger or logging.getLogger(__name__)

    def extract_layout(
        self,
        *,
        docx_path: str | Path,
        template_id: str,
        version: str,
    ) -> LayoutManifest:
        """
        Extract layout metadata from a DOCX file into a LayoutManifest.
        """
        path = Path(docx_path).resolve()

        self._log_info(
            "layout_extraction_start",
            template_id=template_id,
            template_version=version,
            source_docx_path=str(path),
        )

        document = Document(path)
        anchors = self._extract_anchors(document)

        manifest = LayoutManifest(
            template_id=template_id,
            version=version,
            source_docx_path=str(path),
            section_count=len(document.sections),
            anchors=anchors,
            page_setups=self._page_setup_parser.parse(document),
            headers_footers=self._header_footer_parser.parse(document),
            styles=self._style_parser.parse(document),
            tables=self._table_format_parser.parse(document),
        )

        self._log_info(
            "layout_extraction_completed",
            template_id=template_id,
            template_version=version,
            section_count=manifest.section_count,
            anchor_count=len(manifest.anchors),
            style_count=len(manifest.styles),
            table_count=len(manifest.tables),
        )
        return manifest

    def _extract_anchors(self, document: DocxDocument) -> list[AnchorMetadata]:
        """
        Extract heading-like paragraph anchors for later rendering alignment.
        """
        anchors: list[AnchorMetadata] = []
        current_section_index = 0
        max_section_index = max(len(document.sections) - 1, 0)

        for paragraph_index, paragraph in enumerate(document.paragraphs):
            raw_text = paragraph.text.strip()
            if not raw_text:
                self._advance_section_index_if_break(paragraph, current_section_index_ref=[current_section_index], max_section_index=max_section_index)
                current_section_index = self._get_section_index_after_paragraph(paragraph, current_section_index, max_section_index)
                continue

            if not self._is_heading_like(paragraph):
                current_section_index = self._get_section_index_after_paragraph(paragraph, current_section_index, max_section_index)
                continue

            normalized_heading = self._header_normalizer.normalize(raw_text)
            if not normalized_heading:
                current_section_index = self._get_section_index_after_paragraph(paragraph, current_section_index, max_section_index)
                continue

            xml_element_id = paragraph._p.get(self._PARA_ID_ATTR)
            anchor_order = len(anchors)
            anchors.append(
                AnchorMetadata(
                    anchor_id=f"anchor_{anchor_order:04d}",
                    section_index=current_section_index,
                    paragraph_index=paragraph_index,
                    heading_text=raw_text,
                    normalized_heading_text=normalized_heading,
                    xml_element_id=xml_element_id,
                    anchor_order=anchor_order,
                )
            )

            current_section_index = self._get_section_index_after_paragraph(paragraph, current_section_index, max_section_index)

        return anchors

    def _is_heading_like(self, paragraph: Paragraph) -> bool:
        """
        Determine whether a paragraph should be treated as an anchorable heading.
        """
        style_name = (paragraph.style.name or "").strip() if paragraph.style is not None else ""
        if self._HEADING_STYLE_PATTERN.match(style_name):
            return True

        text = paragraph.text.strip()
        numbered_match = self._NUMBERED_HEADING_PATTERN.match(text)
        if numbered_match and self._looks_like_reasonable_fallback_heading(text):
            return True

        return False

    @staticmethod
    def _looks_like_reasonable_fallback_heading(text: str) -> bool:
        """
        Guard the numbered-heading fallback so normal paragraphs are not
        misclassified as headings.
        """
        words = text.split()
        return len(words) <= 15 and len(text) <= 140

    @staticmethod
    def _get_section_index_after_paragraph(
        paragraph: Paragraph,
        current_section_index: int,
        max_section_index: int,
    ) -> int:
        """
        Advance section index if this paragraph ends a Word section.

        python-docx represents section breaks on paragraph properties. We keep
        this logic defensive because not every paragraph exposes the same OXML
        shape in practice.
        """
        p_pr = getattr(paragraph._p, "pPr", None)
        sect_pr = getattr(p_pr, "sectPr", None) if p_pr is not None else None
        if sect_pr is not None and current_section_index < max_section_index:
            return current_section_index + 1
        return current_section_index

    @staticmethod
    def _advance_section_index_if_break(
        paragraph: Paragraph,
        current_section_index_ref: list[int],
        max_section_index: int,
    ) -> None:
        """
        No-op helper retained for readability/possible future enhancement.

        Kept to avoid duplicated conditional branching in `_extract_anchors`.
        """
        current_section_index_ref[:] = current_section_index_ref

    def _log_info(self, event_name: str, **payload: object) -> None:
        """Emit a lightweight structured-ish log entry."""
        self._logger.info("%s | %s", event_name, payload)