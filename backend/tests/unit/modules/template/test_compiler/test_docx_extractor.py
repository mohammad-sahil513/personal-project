"""
Tests for DocxExtractor capturing structural properties.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document

from backend.modules.template.compiler.docx_extractor import DocxExtractor


@pytest.fixture
def test_docx(tmp_path: Path) -> Path:
    """Fixture creating an actual valid minimum .docx file on disk for the test."""
    doc = Document()
    
    # 1. Heading Style
    p_head = doc.add_paragraph("Architecture Overview")
    p_head.style = "Heading 1"
    
    # 2. Numbered Heading Fallback
    doc.add_paragraph("1.2.3 Data Flow")
    
    # 3. Bad numbering fallback test (too long text)
    long_pg = "1. " + ("word " * 30)
    doc.add_paragraph(long_pg)
    
    # 4. Normal paragraph
    doc.add_paragraph("This is just standard content")
    
    # Add a table just to flag contains_table=True
    doc.add_table(rows=1, cols=2)
    
    path = tmp_path / "test.docx"
    doc.save(path)
    return path


def test_docx_extraction_logic(test_docx: Path):
    extractor = DocxExtractor()
    structure = extractor.extract(test_docx)
    
    assert structure.contains_tables is True
    assert structure.contains_multiple_sections is False
    assert structure.contains_headers_footers is False
    
    assert len(structure.headings) == 2
    
    assert structure.headings[0].level == 1
    assert structure.headings[0].raw_text == "Architecture Overview"
    
    assert structure.headings[1].level == 3
    assert structure.headings[1].raw_text == "1.2.3 Data Flow"
