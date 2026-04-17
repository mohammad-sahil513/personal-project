"""
DOCX renderer for assembled workflow documents.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from backend.core.exceptions import ValidationError


class DocxRendererService:
    """
    Render assembled document previews into DOCX files.
    """

    def render(
        self,
        *,
        assembled_document: dict,
        output_path: Path,
    ) -> Path:
        if not assembled_document:
            raise ValidationError(
                message="assembled_document is required for DOCX rendering",
                error_code="DOCX_RENDER_INVALID",
            )

        doc = Document()
        doc.add_heading(assembled_document.get("title", "Document"), level=0)

        for section in assembled_document.get("sections", []):
            doc.add_heading(section["title"], level=1)

            if section.get("content"):
                doc.add_paragraph(section["content"])

            for artifact in section.get("artifacts", []):
                doc.add_paragraph(f"[Artifact: {artifact.get('name')}]")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        return output_path