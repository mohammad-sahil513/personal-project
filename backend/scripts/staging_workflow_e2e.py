#!/usr/bin/env python3
"""
Staging / production HTTP E2E: document upload -> template upload -> compile -> workflow -> output.

Uses the same REST API as the browser. Configure the target base URL and optional file paths via
environment variables (typically alongside your Azure `.env`).

Environment variables
---------------------
STAGING_BASE_URL   Base URL of the running API (no trailing slash). Example: https://api.example.com
                   Default: http://127.0.0.1:8000

E2E_DOCUMENT_PATH  Path to a PDF (or other supported) file to upload for ingestion.
E2E_TEMPLATE_PATH  Path to a .docx template file to upload.

If either path is unset, the script builds minimal placeholder files (may fail real Azure
ingestion/compile — use real paths for a true staging validation).

E2E_POLL_INTERVAL_SEC   Seconds between polls (default: 5)
E2E_MAX_WAIT_SEC        Max seconds to wait for template compile + workflow (default: 900)
E2E_VERIFY_SSL          Set to "false" to disable TLS verification (dev only)

Usage
-----
  cd d:\\ai-sdlc\\backend
  .\\.venv\\Scripts\\python.exe scripts\\staging_workflow_e2e.py

  # Or override URL:
  set STAGING_BASE_URL=https://your-staging-app.azurewebsites.net
  .\\.venv\\Scripts\\python.exe scripts\\staging_workflow_e2e.py

Exit code 0 = all HTTP steps succeeded and output downloaded (or metadata OK if download skipped).
Exit code 1 = failure at any step.
"""

from __future__ import annotations

import io
import os
import sys
import time
from pathlib import Path

from dotenv import load_dotenv

load_dotenv()


def _die(msg: str, code: int = 1) -> None:
    print(f"[E2E][FAIL] {msg}", file=sys.stderr)
    raise SystemExit(code)


def _build_minimal_pdf_bytes() -> bytes:
    """Tiny valid PDF for API-only smoke when E2E_DOCUMENT_PATH is not set."""
    return b"%PDF-1.4\n1 0 obj<<>>endobj\ntrailer<<>>\n%%EOF\n"


def _build_minimal_docx_bytes() -> bytes:
    """Minimal DOCX via python-docx when E2E_TEMPLATE_PATH is not set."""
    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    doc.add_heading("E2E Placeholder Template", level=1)
    doc.add_paragraph("Replace with a real SDLC template for staging.")
    doc.save(buf)
    return buf.getvalue()


def main() -> int:
    try:
        import httpx
    except ImportError as exc:
        _die(f"httpx is required: {exc}")

    base = os.getenv("STAGING_BASE_URL", "http://127.0.0.1:8000").rstrip("/")
    api = f"{base}/api"
    poll_interval = float(os.getenv("E2E_POLL_INTERVAL_SEC", "5"))
    max_wait = float(os.getenv("E2E_MAX_WAIT_SEC", "900"))
    verify_ssl = os.getenv("E2E_VERIFY_SSL", "true").lower() not in ("0", "false", "no")

    doc_path = os.getenv("E2E_DOCUMENT_PATH", "").strip()
    tpl_path = os.getenv("E2E_TEMPLATE_PATH", "").strip()

    timeout = httpx.Timeout(120.0, connect=30.0)
    client = httpx.Client(timeout=timeout, verify=verify_ssl)

    def j(resp: httpx.Response) -> dict:
        if resp.status_code >= 400:
            print(resp.text[:4000], file=sys.stderr)
        resp.raise_for_status()
        payload = resp.json()
        if not payload.get("success", True):
            _die(f"API success=false: {payload}")
        return payload.get("data", payload)

    print(f"[E2E] Base URL: {base}")
    print(f"[E2E] API prefix: {api}")

    # 0) Health
    try:
        h = client.get(f"{api}/health")
        h.raise_for_status()
        health = h.json()
        if not health.get("success", True):
            _die(f"Health check failed: {health}")
        print(f"[E2E] Health: {health.get('data', health)}")
    except httpx.RequestError as exc:
        _die(f"Cannot reach {api}/health: {exc}")

    # 1) Document bytes
    if doc_path:
        p = Path(doc_path)
        if not p.is_file():
            _die(f"E2E_DOCUMENT_PATH not a file: {doc_path}")
        doc_name = p.name
        doc_bytes = p.read_bytes()
        doc_ct = "application/pdf" if doc_name.lower().endswith(".pdf") else "application/octet-stream"
        print(f"[E2E] Using document file: {p} ({len(doc_bytes)} bytes)")
    else:
        doc_bytes = _build_minimal_pdf_bytes()
        doc_name = "e2e_minimal.pdf"
        doc_ct = "application/pdf"
        print("[E2E] E2E_DOCUMENT_PATH unset — using minimal PDF (real ingestion may fail).")

    # 2) Template bytes
    if tpl_path:
        p = Path(tpl_path)
        if not p.is_file():
            _die(f"E2E_TEMPLATE_PATH not a file: {tpl_path}")
        tpl_name = p.name
        tpl_bytes = p.read_bytes()
        print(f"[E2E] Using template file: {p} ({len(tpl_bytes)} bytes)")
    else:
        tpl_bytes = _build_minimal_docx_bytes()
        tpl_name = "e2e_minimal.docx"
        print("[E2E] E2E_TEMPLATE_PATH unset — using minimal DOCX (compile may fail).")

    # 3) Upload document
    files_doc = {"file": (doc_name, doc_bytes, doc_ct)}
    r = client.post(f"{api}/documents/upload", files=files_doc)
    doc_data = j(r)
    document_id = doc_data["document_id"]
    print(f"[E2E] Document uploaded: {document_id}")

    # 4) Upload template
    files_tpl = {
        "file": (tpl_name, tpl_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    }
    data_tpl = {
        "template_type": os.getenv("E2E_TEMPLATE_TYPE", "SDLC"),
        "version": os.getenv("E2E_TEMPLATE_VERSION", "1.0.0"),
    }
    r = client.post(f"{api}/templates/upload", files=files_tpl, data=data_tpl)
    tpl_data = j(r)
    template_id = tpl_data["template_id"]
    print(f"[E2E] Template uploaded: {template_id} status={tpl_data.get('status')}")

    # 5) Start compile
    r = client.post(
        f"{api}/templates/{template_id}/compile",
        json={"use_ai_assist": True, "publish_artifacts": False},
    )
    compile_dispatch = j(r)
    print(f"[E2E] Compile dispatched: {compile_dispatch}")

    # 6) Poll template until COMPILED or terminal failure
    deadline = time.monotonic() + max_wait
    template_status = ""
    while time.monotonic() < deadline:
        r = client.get(f"{api}/templates/{template_id}/compile-status")
        body = j(r)
        template_status = str(body.get("status", ""))
        print(f"[E2E] Template status: {template_status}")
        if template_status == "COMPILED":
            break
        if template_status in ("FAILED",):
            _die(f"Template compile failed (status={template_status})")
        time.sleep(poll_interval)
    else:
        _die(f"Timeout waiting for template COMPILED (last status={template_status})")

    # 7) Create workflow and start
    r = client.post(
        f"{api}/workflow-runs",
        json={
            "document_id": document_id,
            "template_id": template_id,
            "start_immediately": True,
        },
    )
    wf_data = j(r)
    workflow_run_id = wf_data["workflow_run_id"]
    print(f"[E2E] Workflow created: {workflow_run_id} dispatch={wf_data.get('dispatch_mode')}")

    # 8) Poll workflow until COMPLETED or FAILED
    deadline = time.monotonic() + max_wait
    wf_status = ""
    output_id = None
    while time.monotonic() < deadline:
        r = client.get(f"{api}/workflow-runs/{workflow_run_id}")
        body = j(r)
        wf_status = str(body.get("status", ""))
        output_id = body.get("output_id")
        phase = body.get("current_phase", "")
        pct = body.get("overall_progress_percent", 0)
        print(f"[E2E] Workflow status={wf_status} phase={phase} progress={pct}% output_id={output_id}")
        if wf_status == "COMPLETED":
            break
        if wf_status == "FAILED":
            errs = body.get("errors", [])
            _die(f"Workflow FAILED: {errs}")
        time.sleep(poll_interval)
    else:
        _die(f"Timeout waiting for workflow COMPLETED (last status={wf_status})")

    # 9) Download output
    if not output_id:
        _die("Workflow COMPLETED but output_id is missing — check export step.")

    r = client.get(f"{api}/outputs/{output_id}")
    meta = j(r)
    print(f"[E2E] Output metadata: status={meta.get('status')} path={meta.get('artifact_path')}")

    out_path = Path(os.getenv("E2E_OUTPUT_PATH", "e2e_staging_output.docx"))
    r = client.get(f"{api}/outputs/{output_id}/download", follow_redirects=True)
    if r.status_code == 409:
        _die(f"Output not ready for download: {r.text[:500]}")
    r.raise_for_status()
    out_path.write_bytes(r.content)
    print(f"[E2E] Downloaded artifact -> {out_path.resolve()} ({len(r.content)} bytes)")

    print("[E2E] DONE — staging workflow E2E succeeded.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
