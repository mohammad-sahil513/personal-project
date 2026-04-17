"""
Generate a high-fidelity placeholder DOCX using:
- template.json
- layout_manifest.json
- shell.docx
- compiled_artifact.json (optional but recommended)

Goals:
- preserve the authored document flow as much as possible,
- skip TOC/front-matter anchors,
- deduplicate repeated logical sections,
- insert one placeholder block per real body section,
- insert table placeholders only for tabular sections,
- insert image placeholders only for explicit visual/diagram sections.

Preferred source strategy:
1. Use layout_manifest.source_docx_path if it exists (best fidelity).
2. Otherwise fall back to shell.docx.
"""

from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph


PARA_ID_ATTR = "{http://schemas.microsoft.com/office/word/2010/wordml}paraId"

TEXT_PLACEHOLDER_TEMPLATE = """\
{{{{SECTION_ID:{section_id}}}}}
{{{{SECTION_TITLE:{title}}}}}
{{{{TEXT_CONTENT_TO_BE_GENERATED}}}}
"""

IMAGE_PLACEHOLDER_TEMPLATE = """\
{{{{IMAGE_PLACEHOLDER:{section_id}}}}}
{{{{IMAGE_CAPTION:{title}}}}}
{{{{IMAGE_DESCRIPTION_TO_BE_GENERATED}}}}
"""


@dataclass(frozen=True, slots=True)
class SectionInfo:
    """Normalized section metadata derived from template.json."""

    section_id: str
    title: str
    normalized_title: str
    generation_strategy: str | None
    original_index: int


def load_json(path: Path) -> dict[str, Any]:
    """Load and return a JSON object from disk."""
    if not path.exists():
        raise FileNotFoundError(f"Missing JSON file: {path}")
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError(f"JSON root must be an object: {path}")
    return payload


def normalize_heading(text: str) -> str:
    """
    Normalize heading text for matching.

    Handles:
    - leading numbering like '5.' or '5.1'
    - punctuation/tabs
    - trailing page numbers from TOC-like entries
    """
    if not isinstance(text, str):
        return ""

    normalized = text.strip().lower()
    normalized = re.sub(r"^\s*(?:section\s+)?\d+(?:\.\d+)*[\)\.\-:]?\s+", "", normalized)
    normalized = normalized.replace("&", " and ")
    normalized = normalized.replace("/", " ")
    normalized = normalized.replace("_", " ")
    normalized = normalized.replace("-", " ")
    normalized = re.sub(r"[\t\r\n]+", " ", normalized)
    normalized = re.sub(r"[^a-z0-9\s]+", " ", normalized)
    normalized = re.sub(r"\s+", " ", normalized).strip()

    # Remove a trailing page number from TOC lines like "system design 10"
    normalized = re.sub(r"\s+\d+$", "", normalized).strip()
    return normalized


def looks_like_toc_entry(text: str) -> bool:
    """
    Detect likely TOC/list-of-figures/list-of-tables entries.

    Examples:
    - '5. System Design 10'
    - tab-separated heading + page number
    - dot leaders + page number
    """
    raw = text.strip()

    if "\t" in raw and re.search(r"\d+\s*$", raw):
        return True
    if re.search(r"\.{3,}\s*\d+\s*$", raw):
        return True
    if re.search(r"^\s*\d+(?:\.\d+)*[\)\.\-]?\s+.+\s+\d+\s*$", raw):
        return True

    return False


def paragraph_xml_id(paragraph: Paragraph) -> str | None:
    """Return paragraph XML id if present."""
    xml_id = paragraph._p.get(PARA_ID_ATTR)
    if isinstance(xml_id, str) and xml_id.strip():
        return xml_id.strip().upper()
    return None


def build_paragraph_maps(
    document: Document,
) -> tuple[dict[str, Paragraph], dict[str, int], list[Paragraph]]:
    """
    Build paragraph lookup structures for stable anchor resolution.

    Returns:
    - xml id -> paragraph
    - xml id -> paragraph index
    - ordered paragraph list

    Why both maps exist:
    python-docx may return different Paragraph wrapper objects for the same
    underlying XML paragraph, so we should not rely on `paragraphs.index(...)`
    with wrapper identity. Instead, we store the paragraph index directly
    against the stable XML id.
    """
    xml_map: dict[str, Paragraph] = {}
    xml_index_map: dict[str, int] = {}
    paragraphs = list(document.paragraphs)

    for index, paragraph in enumerate(paragraphs):
        xml_id = paragraph_xml_id(paragraph)
        if xml_id:
            xml_map[xml_id] = paragraph
            xml_index_map[xml_id] = index

    return xml_map, xml_index_map, paragraphs


def build_section_infos(template_payload: dict[str, Any]) -> list[SectionInfo]:
    """
    Build normalized SectionInfo entries from template.json.
    """
    sections = template_payload.get("sections", [])
    if not isinstance(sections, list) or not sections:
        raise ValueError("template.json must contain a non-empty `sections` list.")

    result: list[SectionInfo] = []
    for index, section in enumerate(sections):
        if not isinstance(section, dict):
            continue

        section_id = section.get("section_id")
        title = section.get("title")
        generation_strategy = section.get("generation_strategy")

        if not isinstance(section_id, str) or not section_id.strip():
            continue
        if not isinstance(title, str) or not title.strip():
            continue

        result.append(
            SectionInfo(
                section_id=section_id.strip(),
                title=title.strip(),
                normalized_title=normalize_heading(title),
                generation_strategy=generation_strategy if isinstance(generation_strategy, str) else None,
                original_index=index,
            )
        )

    if not result:
        raise ValueError("No usable sections found in template.json.")
    return result


def deduplicate_sections(sections: list[SectionInfo]) -> list[SectionInfo]:
    """
    Collapse repeated logical sections into one high-fidelity placeholder target.

    Deduplication key:
    - normalized title primarily, because the current compiled template may
      contain repeated variants like sec_architecture_description_2, _3, etc.

    Selection rule:
    - prefer the first encountered section (lowest original index),
    - if titles normalize to the same value, keep the earliest one.
    """
    seen: dict[str, SectionInfo] = {}

    for section in sections:
        key = section.normalized_title
        if key not in seen:
            seen[key] = section
            continue

        # Keep the earliest section to preserve authored document order.
        if section.original_index < seen[key].original_index:
            seen[key] = section

    deduplicated = list(seen.values())
    deduplicated.sort(key=lambda item: item.original_index)
    return deduplicated


def validate_artifact_consistency(
    *,
    template_payload: dict[str, Any],
    layout_manifest_payload: dict[str, Any],
    compiled_artifact_payload: dict[str, Any] | None,
) -> None:
    """
    Validate minimal consistency across template/layout/compiled artifacts.
    """
    template_meta = template_payload.get("metadata", {})
    template_id = template_meta.get("template_id")
    template_version = template_meta.get("version")

    layout_template_id = layout_manifest_payload.get("template_id")
    layout_version = layout_manifest_payload.get("version")

    if template_id != layout_template_id:
        raise ValueError(
            f"template.json template_id `{template_id}` does not match layout manifest `{layout_template_id}`."
        )
    if template_version != layout_version:
        raise ValueError(
            f"template.json version `{template_version}` does not match layout manifest `{layout_version}`."
        )

    if compiled_artifact_payload is None:
        return

    artifact_template_id = compiled_artifact_payload.get("template_id")
    artifact_version = compiled_artifact_payload.get("version")

    if template_id != artifact_template_id:
        raise ValueError(
            f"template.json template_id `{template_id}` does not match compiled_artifact `{artifact_template_id}`."
        )
    if template_version != artifact_version:
        raise ValueError(
            f"template.json version `{template_version}` does not match compiled_artifact `{artifact_version}`."
        )


def select_base_docx(
    *,
    layout_manifest_payload: dict[str, Any],
    shell_docx: Path,
) -> Path:
    """
    Prefer the original source DOCX for highest fidelity if it still exists.

    Fallback:
    - shell.docx if source_docx_path is not available.
    """
    source_docx_path = layout_manifest_payload.get("source_docx_path")
    if isinstance(source_docx_path, str) and source_docx_path.strip():
        candidate = Path(source_docx_path).expanduser().resolve()
        if candidate.exists():
            return candidate

    return shell_docx.resolve()


def find_toc_end_index(
    *,
    paragraphs: list[Paragraph],
    normalized_section_titles: set[str],
) -> int:
    """
    Find the index after which real body content likely starts.

    Heuristic:
    - once we see 'table of contents' or a clear TOC/list block,
    - continue scanning until the first non-TOC paragraph that matches a real
      section title from template.json.
    """
    saw_toc = False

    for index, paragraph in enumerate(paragraphs):
        raw_text = paragraph.text.strip()
        normalized = normalize_heading(raw_text)

        if not raw_text:
            continue

        if normalized in {"table of contents", "list of figures", "list of tables"}:
            saw_toc = True
            continue

        if saw_toc:
            if looks_like_toc_entry(raw_text):
                continue
            if normalized in normalized_section_titles:
                return index

    return 0


def match_anchor_candidates_for_section(
    *,
    section: SectionInfo,
    anchors: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """
    Return candidate anchors for one section, scored by normalized title match.
    """
    candidates: list[dict[str, Any]] = []

    for anchor in anchors:
        heading_text = str(anchor.get("heading_text", "")).strip()
        normalized_anchor = normalize_heading(anchor.get("normalized_heading_text") or heading_text)

        if normalized_anchor == section.normalized_title:
            candidates.append(anchor)
            continue

        if section.normalized_title and (
            section.normalized_title in normalized_anchor
            or normalized_anchor in section.normalized_title
        ):
            candidates.append(anchor)

    return candidates


def resolve_body_heading_paragraph(
    *,
    section: SectionInfo,
    anchors: list[dict[str, Any]],
    xml_map: dict[str, Paragraph],
    xml_index_map: dict[str, int],
    paragraphs: list[Paragraph],
    body_start_index: int,
) -> Paragraph | None:
    """
    Resolve the real body heading paragraph for one logical section.

    Resolution order:
    1. XML-level anchor paragraph, but only if it is not TOC-like and lies in body.
    2. paragraph_index anchor, but only if it is not TOC-like and lies in body.
    3. body paragraph text normalized-title match.
    """
    candidate_anchors = match_anchor_candidates_for_section(section=section, anchors=anchors)

    # Prefer anchors with larger paragraph_index because TOC copies generally
    # appear earlier than the real body headings.
    candidate_anchors.sort(
        key=lambda item: int(item.get("paragraph_index", -1))
        if isinstance(item.get("paragraph_index"), int)
        else -1,
        reverse=True,
    )

    for anchor in candidate_anchors:
        xml_id = anchor.get("xml_element_id")
        if isinstance(xml_id, str) and xml_id.strip():
            normalized_xml_id = xml_id.strip().upper()
            paragraph = xml_map.get(normalized_xml_id)
            para_index = xml_index_map.get(normalized_xml_id)

            if paragraph is not None and para_index is not None:
                if para_index >= body_start_index and not looks_like_toc_entry(paragraph.text):
                    return paragraph

        paragraph_index = anchor.get("paragraph_index")
        if isinstance(paragraph_index, int) and 0 <= paragraph_index < len(paragraphs):
            paragraph = paragraphs[paragraph_index]
            if paragraph_index >= body_start_index and not looks_like_toc_entry(paragraph.text):
                return paragraph

    for index, paragraph in enumerate(paragraphs):
        if index < body_start_index:
            continue

        raw_text = paragraph.text.strip()
        if not raw_text:
            continue
        if looks_like_toc_entry(raw_text):
            continue

        normalized = normalize_heading(raw_text)
        if normalized == section.normalized_title:
            return paragraph

    return None


def get_available_table_width(paragraph: Paragraph) -> int:
    """
    Determine a safe available width for inserted placeholder tables.
    """
    try:
        document = paragraph.part.document
        section = document.sections[0]
        page_width = int(section.page_width or 0)
        left_margin = int(section.left_margin or 0)
        right_margin = int(section.right_margin or 0)

        available_width = page_width - left_margin - right_margin
        if available_width > 0:
            return available_width
    except Exception:
        pass

    return int(Inches(6.0))


def insert_paragraph_after(paragraph: Paragraph, text: str, *, style_name: str | None = None) -> Paragraph:
    """
    Insert a new paragraph directly after the given paragraph using XML-level placement.
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)

    if style_name:
        try:
            new_para.style = style_name
        except Exception:
            pass

    if text:
        new_para.add_run(text)

    return new_para


def insert_paragraph_after_table(table: Table, text: str, *, style_name: str | None = None) -> Paragraph:
    """
    Insert a new paragraph directly after the given table using XML-level placement.
    """
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    new_para = Paragraph(new_p, table._parent)

    if style_name:
        try:
            new_para.style = style_name
        except Exception:
            pass

    if text:
        new_para.add_run(text)

    return new_para


def insert_table_after(paragraph: Paragraph, *, style_name: str | None = "Table Grid") -> Table:
    """
    Insert a 3-column placeholder table directly after the given paragraph.
    """
    parent = paragraph._parent
    table_width = get_available_table_width(paragraph)

    table = parent.add_table(rows=2, cols=3, width=table_width)
    paragraph._p.addnext(table._tbl)

    if style_name:
        try:
            table.style = style_name
        except Exception:
            pass

    table.cell(0, 0).text = "{{COLUMN_1}}"
    table.cell(0, 1).text = "{{COLUMN_2}}"
    table.cell(0, 2).text = "{{COLUMN_3}}"

    table.cell(1, 0).text = "{{ROW_1_COL_1}}"
    table.cell(1, 1).text = "{{ROW_1_COL_2}}"
    table.cell(1, 2).text = "{{ROW_1_COL_3}}"

    return table


def needs_table_placeholder(section: SectionInfo) -> bool:
    """
    Insert table placeholders only for explicitly tabular sections.
    """
    return (section.generation_strategy or "").strip().lower() == "generate_table"


def needs_image_placeholder(section: SectionInfo) -> bool:
    """
    Insert image placeholders only for explicitly visual/diagram sections.

    High-fidelity rule:
    - diagram strategy, or
    - heading explicitly contains 'diagram'
    """
    strategy = (section.generation_strategy or "").strip().lower()
    if strategy == "diagram_plantuml":
        return True

    return "diagram" in section.normalized_title


def append_placeholder_to_end(document: Document, section: SectionInfo) -> None:
    """
    Safe fallback when no real body heading can be resolved.
    """
    paragraph = document.add_paragraph(
        TEXT_PLACEHOLDER_TEMPLATE.format(
            section_id=section.section_id,
            title=section.title,
        )
    )

    if needs_table_placeholder(section):
        table = document.add_table(rows=2, cols=3)
        try:
            table.style = "Table Grid"
        except Exception:
            pass

        table.cell(0, 0).text = "{{COLUMN_1}}"
        table.cell(0, 1).text = "{{COLUMN_2}}"
        table.cell(0, 2).text = "{{COLUMN_3}}"
        table.cell(1, 0).text = "{{ROW_1_COL_1}}"
        table.cell(1, 1).text = "{{ROW_1_COL_2}}"
        table.cell(1, 2).text = "{{ROW_1_COL_3}}"

    if needs_image_placeholder(section):
        document.add_paragraph(
            IMAGE_PLACEHOLDER_TEMPLATE.format(
                section_id=section.section_id,
                title=section.title,
            )
        )


def generate_placeholder_docx_high_fidelity(
    *,
    template_json: Path,
    layout_manifest_json: Path,
    shell_docx: Path,
    output_docx: Path,
    compiled_artifact_json: Path | None = None,
) -> None:
    """
    Generate a high-fidelity placeholder DOCX.
    """
    template_payload = load_json(template_json)
    layout_manifest_payload = load_json(layout_manifest_json)
    compiled_artifact_payload = load_json(compiled_artifact_json) if compiled_artifact_json else None

    validate_artifact_consistency(
        template_payload=template_payload,
        layout_manifest_payload=layout_manifest_payload,
        compiled_artifact_payload=compiled_artifact_payload,
    )

    all_sections = build_section_infos(template_payload)
    unique_sections = deduplicate_sections(all_sections)

    base_docx = select_base_docx(
        layout_manifest_payload=layout_manifest_payload,
        shell_docx=shell_docx,
    )
    document = Document(base_docx)

    xml_map, xml_index_map, paragraphs = build_paragraph_maps(document)
    anchors = layout_manifest_payload.get("anchors", [])
    if not isinstance(anchors, list):
        raise ValueError("layout_manifest.json must contain an `anchors` list.")

    normalized_section_titles = {section.normalized_title for section in unique_sections}
    body_start_index = find_toc_end_index(
        paragraphs=paragraphs,
        normalized_section_titles=normalized_section_titles,
    )

    for section in unique_sections:
        anchor_paragraph = resolve_body_heading_paragraph(
            section=section,
            anchors=anchors,
            xml_map=xml_map,
            xml_index_map=xml_index_map,
            paragraphs=paragraphs,
            body_start_index=body_start_index,
        )

        if anchor_paragraph is None:
            append_placeholder_to_end(document, section)
            continue

        style_name = getattr(anchor_paragraph.style, "name", None)

        inserted_text = insert_paragraph_after(
            anchor_paragraph,
            TEXT_PLACEHOLDER_TEMPLATE.format(
                section_id=section.section_id,
                title=section.title,
            ),
            style_name=style_name,
        )

        last_block: Paragraph | Table = inserted_text

        if needs_table_placeholder(section):
            last_block = insert_table_after(inserted_text, style_name="Table Grid")

        if needs_image_placeholder(section):
            if isinstance(last_block, Table):
                insert_paragraph_after_table(
                    last_block,
                    IMAGE_PLACEHOLDER_TEMPLATE.format(
                        section_id=section.section_id,
                        title=section.title,
                    ),
                    style_name=style_name,
                )
            else:
                insert_paragraph_after(
                    last_block,
                    IMAGE_PLACEHOLDER_TEMPLATE.format(
                        section_id=section.section_id,
                        title=section.title,
                    ),
                    style_name=style_name,
                )

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Generate a high-fidelity placeholder DOCX.")
    parser.add_argument("--template-json", required=True, help="Path to template.json")
    parser.add_argument("--layout-manifest-json", required=True, help="Path to layout_manifest.json")
    parser.add_argument("--shell-docx", required=True, help="Path to shell.docx")
    parser.add_argument("--output-docx", required=True, help="Path to output placeholder DOCX")
    parser.add_argument(
        "--compiled-artifact-json",
        default=None,
        help="Optional path to compiled_artifact.json",
    )
    return parser


def main() -> int:
    args = build_arg_parser().parse_args()

    generate_placeholder_docx_high_fidelity(
        template_json=Path(args.template_json).resolve(),
        layout_manifest_json=Path(args.layout_manifest_json).resolve(),
        shell_docx=Path(args.shell_docx).resolve(),
        output_docx=Path(args.output_docx).resolve(),
        compiled_artifact_json=Path(args.compiled_artifact_json).resolve()
        if args.compiled_artifact_json
        else None,
    )

    print("✅ High-fidelity placeholder DOCX generated successfully.")
    print(f"Output: {Path(args.output_docx).resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
