"""
Ingestion smoke test

Checks:
1. Required environment variables exist
2. Required Python packages import correctly
3. Azure Blob Storage connectivity and shared container access
4. Azure AI Document Intelligence connectivity using a generated DOCX
5. Azure AI Search connectivity and index visibility
6. Azure OpenAI chat deployment connectivity
7. Optional embedding deployment connectivity (if configured)

Exit code:
0 -> all required checks passed
1 -> one or more required checks failed
"""

from __future__ import annotations

import io
import json
import os
import sys
import tempfile
import traceback
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional
from dotenv import load_dotenv
load_dotenv()


@dataclass
class CheckResult:
    name: str
    passed: bool
    details: str
    optional: bool = False


@dataclass
class SmokeSummary:
    results: List[CheckResult] = field(default_factory=list)

    def add(self, result: CheckResult) -> None:
        self.results.append(result)

    @property
    def failed_required(self) -> List[CheckResult]:
        return [r for r in self.results if not r.passed and not r.optional]

    @property
    def failed_optional(self) -> List[CheckResult]:
        return [r for r in self.results if not r.passed and r.optional]

    def print_report(self) -> None:
        print("\n===== INGESTION SMOKE TEST REPORT =====\n")
        for result in self.results:
            status = "PASS" if result.passed else ("WARN" if result.optional else "FAIL")
            prefix = f"[{status}] {result.name}"
            print(prefix)
            print(f"      {result.details}")
        print("\n---------------------------------------")
        print(f"Required failures : {len(self.failed_required)}")
        print(f"Optional failures : {len(self.failed_optional)}")
        print(f"Total checks      : {len(self.results)}")
        print("=======================================\n")


def safe_exception_message(exc: BaseException) -> str:
    return f"{type(exc).__name__}: {exc}"


def required_env(name: str) -> str:
    value = os.getenv(name, "").strip()
    if not value:
        raise ValueError(f"Missing required environment variable: {name}")
    return value


def optional_env(name: str) -> str:
    return os.getenv(name, "").strip()


def document_intelligence_endpoint() -> str:
    """Align with .env.example (AZURE_DOCUMENT_INTELLIGENCE_*) and legacy short names."""
    return optional_env("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT") or optional_env(
        "AZURE_DOC_INTELLIGENCE_ENDPOINT"
    )


def document_intelligence_key() -> str:
    return optional_env("AZURE_DOCUMENT_INTELLIGENCE_KEY") or optional_env("AZURE_DOC_INTELLIGENCE_KEY")


def storage_container_name() -> str:
    """Align with .env.example (AZURE_STORAGE_CONTAINER_NAME) and legacy AZURE_STORAGE_CONTAINER."""
    return optional_env("AZURE_STORAGE_CONTAINER_NAME") or optional_env("AZURE_STORAGE_CONTAINER")


def check_env(summary: SmokeSummary) -> None:
    missing: List[str] = []
    if not optional_env("AZURE_OPENAI_ENDPOINT"):
        missing.append("AZURE_OPENAI_ENDPOINT")
    if not optional_env("AZURE_OPENAI_API_KEY"):
        missing.append("AZURE_OPENAI_API_KEY")
    if not optional_env("AZURE_OPENAI_API_VERSION"):
        missing.append("AZURE_OPENAI_API_VERSION")
    if not optional_env("AZURE_OPENAI_REASONING_DEPLOYMENT"):
        missing.append("AZURE_OPENAI_REASONING_DEPLOYMENT")
    if not document_intelligence_endpoint():
        missing.append(
            "AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT (or legacy AZURE_DOC_INTELLIGENCE_ENDPOINT)"
        )
    if not document_intelligence_key():
        missing.append("AZURE_DOCUMENT_INTELLIGENCE_KEY (or legacy AZURE_DOC_INTELLIGENCE_KEY)")
    if not optional_env("AZURE_SEARCH_ENDPOINT"):
        missing.append("AZURE_SEARCH_ENDPOINT")
    if not optional_env("AZURE_SEARCH_API_KEY"):
        missing.append("AZURE_SEARCH_API_KEY")
    if not optional_env("AZURE_SEARCH_INDEX_NAME"):
        missing.append("AZURE_SEARCH_INDEX_NAME")
    if not optional_env("AZURE_STORAGE_CONNECTION_STRING"):
        missing.append("AZURE_STORAGE_CONNECTION_STRING")
    if not storage_container_name():
        missing.append("AZURE_STORAGE_CONTAINER_NAME (or legacy AZURE_STORAGE_CONTAINER)")
    if not optional_env("AZURE_STORAGE_ROOT_PREFIX"):
        missing.append("AZURE_STORAGE_ROOT_PREFIX")
    if missing:
        summary.add(CheckResult(
            name="Environment variables",
            passed=False,
            details="Missing: " + ", ".join(missing),
        ))
    else:
        summary.add(CheckResult(
            name="Environment variables",
            passed=True,
            details="All required ingestion environment variables are present.",
        ))


def check_imports(summary: SmokeSummary) -> None:
    imports = {
        "openai": "Azure OpenAI / Responses API client",
        "semantic_kernel": "Semantic Kernel",
        "azure.identity": "Azure Identity",
        "azure.search.documents": "Azure AI Search",
        "azure.ai.documentintelligence": "Azure AI Document Intelligence",
        "azure.storage.blob": "Azure Blob Storage",
        "docx": "python-docx",
        "structlog": "structured logging",
        "httpx": "HTTP client",
        "aiohttp": "async transport",
    }
    missing = []
    for module_name in imports:
        try:
            __import__(module_name)
        except Exception as exc:
            missing.append(f"{module_name} ({safe_exception_message(exc)})")
    if missing:
        summary.add(CheckResult(
            name="Python package imports",
            passed=False,
            details="Import failures: " + "; ".join(missing),
        ))
    else:
        summary.add(CheckResult(
            name="Python package imports",
            passed=True,
            details="All required ingestion/runtime packages imported successfully.",
        ))


def check_blob_storage(summary: SmokeSummary) -> None:
    try:
        from azure.storage.blob import BlobServiceClient

        connection_string = required_env("AZURE_STORAGE_CONNECTION_STRING")
        container_name = storage_container_name()
        if not container_name:
            raise ValueError(
                "Missing AZURE_STORAGE_CONTAINER_NAME or AZURE_STORAGE_CONTAINER"
            )
        root_prefix = required_env("AZURE_STORAGE_ROOT_PREFIX")

        service = BlobServiceClient.from_connection_string(connection_string)
        container = service.get_container_client(container_name)
        exists = container.exists()
        if not exists:
            raise RuntimeError(f"Container does not exist or is not accessible: {container_name}")

        # Lightweight probe: list up to 5 blobs under the shared root prefix.
        blobs = list(container.list_blobs(name_starts_with=f"{root_prefix}/"))[:5]
        details = (
            f"Connected to container '{container_name}'. "
            f"Shared root prefix '{root_prefix}/' is accessible. "
            f"Found {len(blobs)} blob(s) under the prefix (showing max 5)."
        )
        summary.add(CheckResult("Azure Blob Storage", True, details))
    except Exception as exc:
        summary.add(CheckResult(
            name="Azure Blob Storage",
            passed=False,
            details=safe_exception_message(exc),
        ))


def build_minimal_docx_bytes() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Smoke Test Document", level=1)
    doc.add_paragraph("This is a generated DOCX used for Azure AI Document Intelligence smoke testing.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def check_document_intelligence(summary: SmokeSummary) -> None:
    try:
        from azure.core.credentials import AzureKeyCredential
        from azure.ai.documentintelligence import DocumentIntelligenceClient

        endpoint = document_intelligence_endpoint()
        key = document_intelligence_key()
        if not endpoint or not key:
            raise ValueError(
                "Missing Document Intelligence endpoint/key "
                "(AZURE_DOCUMENT_INTELLIGENCE_* or legacy AZURE_DOC_INTELLIGENCE_*)"
            )

        client = DocumentIntelligenceClient(endpoint=endpoint, credential=AzureKeyCredential(key))
        content = build_minimal_docx_bytes()
        poller = client.begin_analyze_document(
            model_id="prebuilt-layout",
            body=content,
            output_content_format="markdown",
        )
        result = poller.result()
        pages = getattr(result, "pages", None)
        page_count = len(pages) if pages else 0
        summary.add(CheckResult(
            name="Azure AI Document Intelligence",
            passed=True,
            details=f"prebuilt-layout call succeeded on generated DOCX. Pages returned: {page_count}.",
        ))
    except Exception as exc:
        summary.add(CheckResult(
            name="Azure AI Document Intelligence",
            passed=False,
            details=safe_exception_message(exc),
        ))


def check_search(summary: SmokeSummary) -> None:
    try:
        from azure.core.credentials import AzureKeyCredential
        from azure.search.documents.indexes import SearchIndexClient

        endpoint = required_env("AZURE_SEARCH_ENDPOINT")
        key = required_env("AZURE_SEARCH_API_KEY")
        expected_index = required_env("AZURE_SEARCH_INDEX_NAME")

        client = SearchIndexClient(endpoint=endpoint, credential=AzureKeyCredential(key))
        index_names = [idx.name for idx in client.list_indexes()]
        if expected_index not in index_names:
            raise RuntimeError(
                f"Search service reachable, but expected index '{expected_index}' was not found. "
                f"Existing indexes: {index_names or 'none'}"
            )
        summary.add(CheckResult(
            name="Azure AI Search",
            passed=True,
            details=f"Connected successfully. Expected index '{expected_index}' is present.",
        ))
    except Exception as exc:
        summary.add(CheckResult(
            name="Azure AI Search",
            passed=False,
            details=safe_exception_message(exc),
        ))


def check_openai_chat(summary: SmokeSummary) -> None:
    try:
        from openai import OpenAI

        endpoint = required_env("AZURE_OPENAI_ENDPOINT")
        api_key = required_env("AZURE_OPENAI_API_KEY")
        deployment = required_env("AZURE_OPENAI_REASONING_DEPLOYMENT")

        client = OpenAI(api_key=api_key, base_url=f"{endpoint.rstrip('/')}/openai/v1/")
        response = client.responses.create(
            model=deployment,
            input="Reply with exactly the word: OK",
        )
        text = getattr(response, "output_text", "") or ""
        summary.add(CheckResult(
            name="Azure OpenAI chat deployment",
            passed=True,
            details=f"Chat deployment '{deployment}' responded successfully. Sample output: {text[:60]!r}",
        ))
    except Exception as exc:
        summary.add(CheckResult(
            name="Azure OpenAI chat deployment",
            passed=False,
            details=safe_exception_message(exc),
        ))


def check_openai_embeddings(summary: SmokeSummary) -> None:
    deployment = optional_env("AZURE_OPENAI_EMBEDDING_DEPLOYMENT")
    if not deployment:
        summary.add(CheckResult(
            name="Azure OpenAI embedding deployment",
            passed=False,
            details="Embedding deployment not configured in .env. Skipping optional check.",
            optional=True,
        ))
        return
    try:
        from openai import OpenAI

        endpoint = required_env("AZURE_OPENAI_ENDPOINT")
        api_key = required_env("AZURE_OPENAI_API_KEY")

        client = OpenAI(api_key=api_key, base_url=f"{endpoint.rstrip('/')}/openai/v1/")
        response = client.embeddings.create(
            model=deployment,
            input="smoke test embedding",
        )
        vec_len = len(response.data[0].embedding) if response.data else 0
        summary.add(CheckResult(
            name="Azure OpenAI embedding deployment",
            passed=True,
            details=f"Embedding deployment '{deployment}' responded successfully. Vector length: {vec_len}.",
            optional=True,
        ))
    except Exception as exc:
        summary.add(CheckResult(
            name="Azure OpenAI embedding deployment",
            passed=False,
            details=safe_exception_message(exc),
            optional=True,
        ))


def main() -> int:
    summary = SmokeSummary()
    check_env(summary)
    check_imports(summary)

    # Only continue with service checks if env and imports are okay enough to proceed.
    if not summary.failed_required:
        check_blob_storage(summary)
        check_document_intelligence(summary)
        check_search(summary)
        check_openai_chat(summary)
        check_openai_embeddings(summary)

    summary.print_report()

    if summary.failed_required:
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
