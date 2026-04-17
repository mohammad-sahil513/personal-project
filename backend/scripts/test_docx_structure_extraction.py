"""
Standalone deterministic DOCX structure extraction test runner.

Purpose:
- test custom-template DOCX extraction without any LLM usage,
- extract headings/sub-headings,
- extract tables,
- extract images (including likely logos / inline images),
- extract section headers and footers,
- write JSON artifacts for inspection.

Outputs under the chosen output directory:
- docx_structure.json
- docx_tables.json
- docx_images.json
- docx_headers.json
- docx_assets/ (extracted images)

Usage:
python backend/scripts/test_docx_structure_extraction.py \
  --docx ".\\backend\\data\\sdd_sample1.docx" \
  --output-dir ".\\backend\\artifacts\\template\\docx_extraction_test"
"""

from __future__ import annotations

import argparse
import json
import mimetypes
import re
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.shared import Emu
from docx.table import Table
from docx.text.paragraph import Paragraph


PARA_ID_ATTR = "{http://schemas.microsoft.com/office/word/2010/wordml}paraId"
HEADING_STYLE_PATTERN = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)
NUMBERED_HEADING_PATTERN = re.compile(r"^\s*(\d+(?:\.\d+)*)[\)\.\-]?\s+(.+?)\s*$")
REL_IMAGE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"


@dataclass(frozen=True, slots=True)
class HeadingRecord:
    section_id: str
    level: int
    title: str
    normalized_title: str
    style_name: str | None
    paragraph_index: int
    paragraph_xml_id: str | None
    start_index: int
    end_index: int


@dataclass(frozen=True, slots=True)
class TableRecord:
    table_id: str
    section_id: str | None
    table_index: int
    paragraph_index_hint: int | None
    rows: int
    columns: int
    style_name: str | None
    cells: list[list[str]]


@dataclass(frozen=True, slots=True)
class ImageRecord:
    image_id: str
    section_id: str | None
    paragraph_index: int | None
    relationship_id: str
    filename: str
    content_type: str | None
    width_emu: int | None
    height_emu: int | None
    inferred_kind: str


@dataclass(frozen=True, slots=True)
class HeaderFooterRecord:
    section_index: int
    header_text: str
    footer_text: str
    header_linked_to_previous: bool
    footer_linked_to_previous: bool


class DeterministicDocxExtractor:
    """Deterministic DOCX structure extractor for testing custom-template handling."""

    def __init__(self, *, docx_path: str | Path, output_dir: str | Path) -> None:
        self.docx_path = Path(docx_path).resolve()
        self.output_dir = Path(output_dir).resolve()
        self.assets_dir = self.output_dir / "docx_assets"
        self.document: DocxDocument = Document(self.docx_path)

    def run(self) -> dict[str, Any]:
        """Run full extraction and write output files."""
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.assets_dir.mkdir(parents=True, exist_ok=True)

        headings = self.extract_headings()
        tables = self.extract_tables(headings)
        images = self.extract_images(headings)
        headers = self.extract_headers_footers()

        structure_payload = {
            "document": {
                "source_docx_path": str(self.docx_path),
                "section_count": len(self.document.sections),
                "paragraph_count": len(self.document.paragraphs),
                "table_count": len(tables),
                "image_count": len(images),
            },
            "sections": [asdict(item) for item in headings],
        }
        tables_payload = {"tables": [asdict(item) for item in tables]}
        images_payload = {"images": [asdict(item) for item in images]}
        headers_payload = {"headers_footers": [asdict(item) for item in headers]}

        self._write_json(self.output_dir / "docx_structure.json", structure_payload)
        self._write_json(self.output_dir / "docx_tables.json", tables_payload)
        self._write_json(self.output_dir / "docx_images.json", images_payload)
        self._write_json(self.output_dir / "docx_headers.json", headers_payload)

        return {
            "structure": structure_payload,
            "tables": tables_payload,
            "images": images_payload,
            "headers_footers": headers_payload,
        }

    def extract_headings(self) -> list[HeadingRecord]:
        """Extract headings and derive section ranges."""
        paragraphs = list(self.document.paragraphs)
        provisional: list[dict[str, Any]] = []

        for paragraph_index, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            level = self._infer_heading_level(paragraph)
            if level is None:
                continue

            normalized_title = self.normalize_heading(text)
            if not normalized_title:
                continue

            provisional.append(
                {
                    "level": level,
                    "title": text,
                    "normalized_title": normalized_title,
                    "style_name": getattr(paragraph.style, "name", None),
                    "paragraph_index": paragraph_index,
                    "paragraph_xml_id": paragraph._p.get(PARA_ID_ATTR),
                }
            )

        headings: list[HeadingRecord] = []
        for index, item in enumerate(provisional):
            start_index = item["paragraph_index"]
            end_index = (
                provisional[index + 1]["paragraph_index"] - 1
                if index + 1 < len(provisional)
                else len(paragraphs) - 1
            )

            headings.append(
                HeadingRecord(
                    section_id=self._build_section_id(item["normalized_title"], index),
                    level=item["level"],
                    title=item["title"],
                    normalized_title=item["normalized_title"],
                    style_name=item["style_name"],
                    paragraph_index=item["paragraph_index"],
                    paragraph_xml_id=item["paragraph_xml_id"],
                    start_index=start_index,
                    end_index=end_index,
                )
            )

        return headings

    def extract_tables(self, headings: list[HeadingRecord]) -> list[TableRecord]:
        """Extract tables and map them to nearest section by position."""
        records: list[TableRecord] = []
        paragraph_map = {id(p._p): i for i, p in enumerate(self.document.paragraphs)}

        for table_index, table in enumerate(self.document.tables):
            cells = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            row_count = len(table.rows)
            column_count = max((len(row.cells) for row in table.rows), default=0)
            style_name = table.style.name if table.style is not None else None

            paragraph_index_hint = self._infer_table_paragraph_index(table, paragraph_map)
            section_id = self._resolve_section_for_paragraph_index(headings, paragraph_index_hint)

            records.append(
                TableRecord(
                    table_id=f"tbl_{table_index + 1:03d}",
                    section_id=section_id,
                    table_index=table_index,
                    paragraph_index_hint=paragraph_index_hint,
                    rows=row_count,
                    columns=column_count,
                    style_name=style_name,
                    cells=cells,
                )
            )

        return records

    def extract_images(self, headings: list[HeadingRecord]) -> list[ImageRecord]:
        """Extract images from DOCX relationships and write them to docx_assets/."""
        records: list[ImageRecord] = []
        rels = self.document.part.rels
        drawing_to_paragraph_index = self._build_drawing_paragraph_index_map()

        image_counter = 0
        for rel_id, rel in rels.items():
            if rel.reltype != REL_IMAGE:
                continue

            image_counter += 1
            image_part = rel.target_part
            image_bytes = image_part.blob
            content_type = getattr(image_part, "content_type", None)
            ext = mimetypes.guess_extension(content_type or "") or ".bin"
            filename = f"img_{image_counter:03d}{ext}"
            output_path = self.assets_dir / filename
            output_path.write_bytes(image_bytes)

            paragraph_index = drawing_to_paragraph_index.get(rel_id)
            section_id = self._resolve_section_for_paragraph_index(headings, paragraph_index)
            width_emu, height_emu = self._infer_inline_shape_size(rel_id)

            inferred_kind = self._infer_image_kind(
                paragraph_index=paragraph_index,
                width_emu=width_emu,
                height_emu=height_emu,
            )

            records.append(
                ImageRecord(
                    image_id=f"img_{image_counter:03d}",
                    section_id=section_id,
                    paragraph_index=paragraph_index,
                    relationship_id=rel_id,
                    filename=filename,
                    content_type=content_type,
                    width_emu=width_emu,
                    height_emu=height_emu,
                    inferred_kind=inferred_kind,
                )
            )

        return records

    def extract_headers_footers(self) -> list[HeaderFooterRecord]:
        """Extract per-section headers and footers."""
        records: list[HeaderFooterRecord] = []

        for section_index, section in enumerate(self.document.sections):
            header_text = "\n".join(
                p.text.strip() for p in section.header.paragraphs if p.text and p.text.strip()
            )
            footer_text = "\n".join(
                p.text.strip() for p in section.footer.paragraphs if p.text and p.text.strip()
            )

            records.append(
                HeaderFooterRecord(
                    section_index=section_index,
                    header_text=header_text,
                    footer_text=footer_text,
                    header_linked_to_previous=bool(section.header.is_linked_to_previous),
                    footer_linked_to_previous=bool(section.footer.is_linked_to_previous),
                )
            )

        return records

    @staticmethod
    def normalize_heading(text: str) -> str:
        """Normalize heading text for deterministic matching."""
        normalized = text.strip().lower()
        normalized = re.sub(r"^\s*(?:section\s+)?\d+(?:\.\d+)*[\)\.\-:]?\s+", "", normalized)
        normalized = normalized.replace("&", " and ")
        normalized = normalized.replace("/", " ")
        normalized = normalized.replace("_", " ")
        normalized = normalized.replace("-", " ")
        normalized = re.sub(r"[\t\r\n]+", " ", normalized)
        normalized = re.sub(r"[^a-z0-9\s]+", " ", normalized)
        normalized = re.sub(r"\s+", " ", normalized).strip()
        normalized = re.sub(r"\s+\d+$", "", normalized).strip()
        return normalized

    @staticmethod
    def _write_json(path: Path, payload: dict[str, Any]) -> None:
        """Write JSON file to disk."""
        path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")

    @staticmethod
    def _build_section_id(normalized_title: str, index: int) -> str:
        """Build a stable synthetic section id for testing."""
        slug = normalized_title.replace(" ", "_") or f"section_{index + 1}"
        return f"sec_{slug}_{index + 1}"

    @staticmethod
    def _infer_heading_level(paragraph: Paragraph) -> int | None:
        """Infer heading level from style or bounded numbered-heading fallback."""
        style_name = (paragraph.style.name or "").strip() if paragraph.style is not None else ""
        match = HEADING_STYLE_PATTERN.match(style_name)
        if match:
            return max(1, min(9, int(match.group(1))))

        text = paragraph.text.strip()
        numbered = NUMBERED_HEADING_PATTERN.match(text)
        if numbered and len(text.split()) <= 15 and len(text) <= 140:
            numbering = numbered.group(1)
            return max(1, min(9, numbering.count(".") + 1))

        return None

    @staticmethod
    def _resolve_section_for_paragraph_index(
        headings: list[HeadingRecord],
        paragraph_index: int | None,
    ) -> str | None:
        """Map a paragraph index to the heading range that contains it."""
        if paragraph_index is None:
            return None

        for heading in headings:
            if heading.start_index <= paragraph_index <= heading.end_index:
                return heading.section_id

        return None

    @staticmethod
    def _infer_table_paragraph_index(table: Table, paragraph_map: dict[int, int]) -> int | None:
        """
        Approximate the table's position by finding the nearest preceding paragraph.
        """
        try:
            previous = table._tbl.getprevious()
            while previous is not None:
                if previous.tag == qn("w:p"):
                    return paragraph_map.get(id(previous))
                previous = previous.getprevious()
        except Exception:
            pass
        return None

    def _build_drawing_paragraph_index_map(self) -> dict[str, int]:
        """
        Map image relationship ids to paragraph indices by scanning drawing blips.
        """
        mapping: dict[str, int] = {}

        for paragraph_index, paragraph in enumerate(self.document.paragraphs):
            try:
                blips = paragraph._p.xpath(".//a:blip")
            except Exception:
                blips = []

            for blip in blips:
                embed = blip.get(qn("r:embed"))
                if isinstance(embed, str) and embed.strip() and embed not in mapping:
                    mapping[embed] = paragraph_index

        return mapping

    def _infer_inline_shape_size(self, rel_id: str) -> tuple[int | None, int | None]:
        """
        Find inline shape width/height for a given image relationship id.
        """
        for shape in self.document.inline_shapes:
            try:
                blips = shape._inline.xpath(".//a:blip")
            except Exception:
                blips = []

            for blip in blips:
                embed = blip.get(qn("r:embed"))
                if embed == rel_id:
                    width = int(shape.width) if isinstance(shape.width, Emu) or shape.width is not None else None
                    height = int(shape.height) if isinstance(shape.height, Emu) or shape.height is not None else None
                    return width, height

        return None, None

    @staticmethod
    def _infer_image_kind(
        *,
        paragraph_index: int | None,
        width_emu: int | None,
        height_emu: int | None,
    ) -> str:
        """
        Very conservative deterministic guess only for test visibility.
        """
        if paragraph_index is not None and paragraph_index < 10:
            return "logo"
        if width_emu and height_emu and width_emu < 2_500_000 and height_emu < 1_500_000:
            return "icon_or_logo"
        return "inline_image"


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Run deterministic DOCX extraction test.")
    parser.add_argument("--docx", required=True, help="Path to the source DOCX file.")
    parser.add_argument(
        "--output-dir",
        required=True,
        help="Directory where extracted JSON and assets should be written.",
    )
    return parser


def main() -> int:
    args = build_arg_parser().parse_args()

    extractor = DeterministicDocxExtractor(
        docx_path=args.docx,
        output_dir=args.output_dir,
    )
    result = extractor.run()

    summary = {
        "status": "success",
        "docx": str(Path(args.docx).resolve()),
        "output_dir": str(Path(args.output_dir).resolve()),
        "section_count": len(result["structure"]["sections"]),
        "table_count": len(result["tables"]["tables"]),
        "image_count": len(result["images"]["images"]),
        "header_footer_count": len(result["headers_footers"]["headers_footers"]),
    }
    print(json.dumps(summary, indent=2, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())