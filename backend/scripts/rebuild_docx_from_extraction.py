"""
Rebuild a structural DOCX from deterministic extraction artifacts.

This script uses:
- docx_structure.json
- docx_tables.json
- docx_images.json
- docx_headers.json
- docx_assets/

to generate a new DOCX containing:
- document headers / footers
- heading hierarchy
- extracted tables
- extracted images / logos
- explicit placeholder body text for sections whose narrative body text was not
  captured during deterministic extraction

Important limitation:
This rebuild is a structural reconstruction, not a byte-faithful clone of the
original DOCX. To achieve full-fidelity round-tripping, the extraction phase
would also need a complete ordered block stream (paragraphs/tables/images) and
section-index mapping for every block.
"""

from __future__ import annotations

import argparse
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Emu


DEFAULT_BODY_PLACEHOLDER = "[Original body content was not captured in extraction artifacts.]"


@dataclass(frozen=True, slots=True)
class SectionRecord:
    """Heading/section record extracted from docx_structure.json."""

    section_id: str
    level: int
    title: str
    paragraph_index: int


@dataclass(frozen=True, slots=True)
class TableRecord:
    """Table record extracted from docx_tables.json."""

    table_id: str
    section_id: str | None
    table_index: int
    paragraph_index_hint: int | None
    rows: int
    columns: int
    style_name: str | None
    cells: list[list[str]]


@dataclass(frozen=True, slots=True)
class ImageRecord:
    """Image record extracted from docx_images.json."""

    image_id: str
    section_id: str | None
    paragraph_index: int | None
    relationship_id: str
    filename: str
    content_type: str | None
    width_emu: int | None
    height_emu: int | None
    inferred_kind: str


@dataclass(frozen=True, slots=True)
class HeaderFooterRecord:
    """Header/footer record extracted from docx_headers.json."""

    section_index: int
    header_text: str
    footer_text: str
    header_linked_to_previous: bool
    footer_linked_to_previous: bool


def _load_json(path: Path) -> dict[str, Any]:
    """Load a JSON file and validate that the root is an object."""
    if not path.exists():
        raise FileNotFoundError(f"Missing JSON file: {path}")

    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError(f"JSON root must be an object: {path}")

    return payload


def _load_sections(path: Path) -> list[SectionRecord]:
    """Load section records from docx_structure.json."""
    payload = _load_json(path)
    raw_sections = payload.get("sections", [])
    if not isinstance(raw_sections, list):
        raise ValueError("docx_structure.json must contain a `sections` list.")

    sections: list[SectionRecord] = []
    for item in raw_sections:
        if not isinstance(item, dict):
            continue

        section_id = item.get("section_id")
        level = item.get("level")
        title = item.get("title")
        paragraph_index = item.get("paragraph_index")

        if not isinstance(section_id, str) or not section_id.strip():
            continue
        if not isinstance(level, int) or level < 1:
            continue
        if not isinstance(title, str) or not title.strip():
            continue
        if not isinstance(paragraph_index, int) or paragraph_index < 0:
            continue

        sections.append(
            SectionRecord(
                section_id=section_id.strip(),
                level=level,
                title=title.strip(),
                paragraph_index=paragraph_index,
            )
        )

    sections.sort(key=lambda item: item.paragraph_index)
    return sections


def _load_tables(path: Path) -> list[TableRecord]:
    """Load table records from docx_tables.json."""
    payload = _load_json(path)
    raw_tables = payload.get("tables", [])
    if not isinstance(raw_tables, list):
        raise ValueError("docx_tables.json must contain a `tables` list.")

    tables: list[TableRecord] = []
    for item in raw_tables:
        if not isinstance(item, dict):
            continue

        cells = item.get("cells", [])
        tables.append(
            TableRecord(
                table_id=str(item.get("table_id", "")).strip(),
                section_id=item.get("section_id") if isinstance(item.get("section_id"), str) else None,
                table_index=int(item.get("table_index", 0)),
                paragraph_index_hint=(
                    item.get("paragraph_index_hint")
                    if isinstance(item.get("paragraph_index_hint"), int)
                    else None
                ),
                rows=int(item.get("rows", 0)),
                columns=int(item.get("columns", 0)),
                style_name=item.get("style_name") if isinstance(item.get("style_name"), str) else None,
                cells=cells if isinstance(cells, list) else [],
            )
        )

    tables.sort(
        key=lambda item: (
            item.paragraph_index_hint if item.paragraph_index_hint is not None else 10**9,
            item.table_index,
        )
    )
    return tables


def _load_images(path: Path) -> list[ImageRecord]:
    """Load image records from docx_images.json."""
    payload = _load_json(path)
    raw_images = payload.get("images", [])
    if not isinstance(raw_images, list):
        raise ValueError("docx_images.json must contain an `images` list.")

    images: list[ImageRecord] = []
    for item in raw_images:
        if not isinstance(item, dict):
            continue

        images.append(
            ImageRecord(
                image_id=str(item.get("image_id", "")).strip(),
                section_id=item.get("section_id") if isinstance(item.get("section_id"), str) else None,
                paragraph_index=item.get("paragraph_index") if isinstance(item.get("paragraph_index"), int) else None,
                relationship_id=str(item.get("relationship_id", "")).strip(),
                filename=str(item.get("filename", "")).strip(),
                content_type=item.get("content_type") if isinstance(item.get("content_type"), str) else None,
                width_emu=item.get("width_emu") if isinstance(item.get("width_emu"), int) else None,
                height_emu=item.get("height_emu") if isinstance(item.get("height_emu"), int) else None,
                inferred_kind=str(item.get("inferred_kind", "inline_image")).strip() or "inline_image",
            )
        )

    images.sort(
        key=lambda item: (
            item.paragraph_index if item.paragraph_index is not None else 10**9,
            item.image_id,
        )
    )
    return images


def _load_headers(path: Path) -> list[HeaderFooterRecord]:
    """Load header/footer records from docx_headers.json."""
    payload = _load_json(path)
    raw_headers = payload.get("headers_footers", [])
    if not isinstance(raw_headers, list):
        raise ValueError("docx_headers.json must contain a `headers_footers` list.")

    headers: list[HeaderFooterRecord] = []
    for item in raw_headers:
        if not isinstance(item, dict):
            continue

        headers.append(
            HeaderFooterRecord(
                section_index=int(item.get("section_index", 0)),
                header_text=item.get("header_text") if isinstance(item.get("header_text"), str) else "",
                footer_text=item.get("footer_text") if isinstance(item.get("footer_text"), str) else "",
                header_linked_to_previous=bool(item.get("header_linked_to_previous", False)),
                footer_linked_to_previous=bool(item.get("footer_linked_to_previous", False)),
            )
        )

    headers.sort(key=lambda item: item.section_index)
    return headers


def _group_tables_by_section(tables: list[TableRecord]) -> dict[str | None, list[TableRecord]]:
    """Group tables by section_id."""
    grouped: dict[str | None, list[TableRecord]] = {}
    for table in tables:
        grouped.setdefault(table.section_id, []).append(table)
    return grouped


def _group_images_by_section(images: list[ImageRecord]) -> dict[str | None, list[ImageRecord]]:
    """Group images by section_id."""
    grouped: dict[str | None, list[ImageRecord]] = {}
    for image in images:
        grouped.setdefault(image.section_id, []).append(image)
    return grouped


def _apply_headers_and_footers(document: Document, headers: list[HeaderFooterRecord]) -> None:
    """Apply first-section header/footer content to the rebuilt document."""
    if not headers:
        return

    first = headers[0]
    document.sections[0].header.paragraphs[0].text = first.header_text
    document.sections[0].footer.paragraphs[0].text = first.footer_text


def _add_image(document: Document, assets_dir: Path, image: ImageRecord) -> None:
    """Insert an extracted image into the rebuilt DOCX."""
    image_path = assets_dir / image.filename
    if not image_path.exists():
        document.add_paragraph(f"[Missing extracted image asset: {image.filename}]")
        return

    kwargs: dict[str, Any] = {}
    if image.width_emu and image.width_emu > 0:
        kwargs["width"] = Emu(image.width_emu)

    document.add_picture(str(image_path), **kwargs)
    document.add_paragraph(f"[Extracted image: {image.filename} | kind={image.inferred_kind}]")


def _add_table(document: Document, table_record: TableRecord) -> None:
    """Insert an extracted table into the rebuilt DOCX."""
    row_count = max(table_record.rows, len(table_record.cells), 1)
    column_count = max(
        table_record.columns,
        max((len(row) for row in table_record.cells), default=0),
        1,
    )

    table = document.add_table(rows=row_count, cols=column_count)
    if table_record.style_name:
        try:
            table.style = table_record.style_name
        except Exception:
            # Best effort only; style may not exist in target document style set.
            pass

    for row_idx in range(row_count):
        source_row = table_record.cells[row_idx] if row_idx < len(table_record.cells) else []
        for col_idx in range(column_count):
            value = source_row[col_idx] if col_idx < len(source_row) else ""
            table.cell(row_idx, col_idx).text = value


def rebuild_docx_from_extraction(
    *,
    docx_structure_json: Path,
    docx_tables_json: Path,
    docx_images_json: Path,
    docx_headers_json: Path,
    docx_assets_dir: Path,
    output_docx: Path,
    include_body_placeholder: bool = True,
) -> None:
    """
    Rebuild a structural DOCX from deterministic extraction artifacts.
    """
    sections = _load_sections(docx_structure_json)
    tables = _load_tables(docx_tables_json)
    images = _load_images(docx_images_json)
    headers = _load_headers(docx_headers_json)

    tables_by_section = _group_tables_by_section(tables)
    images_by_section = _group_images_by_section(images)

    document = Document()
    _apply_headers_and_footers(document, headers)

    # Add front-matter images / logos
    for image in images_by_section.get(None, []):
        _add_image(document, docx_assets_dir, image)

    # Add section headings + placeholder body + mapped images/tables
    for section in sections:
        heading_level = max(1, min(section.level, 9))
        document.add_heading(section.title, level=heading_level)

        if include_body_placeholder:
            document.add_paragraph(DEFAULT_BODY_PLACEHOLDER)

        for image in images_by_section.get(section.section_id, []):
            _add_image(document, docx_assets_dir, image)

        for table_record in tables_by_section.get(section.section_id, []):
            _add_table(document, table_record)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Rebuild DOCX from deterministic extraction artifacts.")
    parser.add_argument("--docx-structure-json", required=True, help="Path to docx_structure.json")
    parser.add_argument("--docx-tables-json", required=True, help="Path to docx_tables.json")
    parser.add_argument("--docx-images-json", required=True, help="Path to docx_images.json")
    parser.add_argument("--docx-headers-json", required=True, help="Path to docx_headers.json")
    parser.add_argument("--docx-assets-dir", required=True, help="Path to docx_assets directory")
    parser.add_argument("--output-docx", required=True, help="Path to output reconstructed DOCX")
    parser.add_argument(
        "--no-body-placeholder",
        action="store_true",
        help="Do not insert placeholder body paragraph under each heading.",
    )
    return parser


def main() -> int:
    args = build_arg_parser().parse_args()

    rebuild_docx_from_extraction(
        docx_structure_json=Path(args.docx_structure_json).resolve(),
        docx_tables_json=Path(args.docx_tables_json).resolve(),
        docx_images_json=Path(args.docx_images_json).resolve(),
        docx_headers_json=Path(args.docx_headers_json).resolve(),
        docx_assets_dir=Path(args.docx_assets_dir).resolve(),
        output_docx=Path(args.output_docx).resolve(),
        include_body_placeholder=not args.no_body_placeholder,
    )

    print("✅ Reconstructed DOCX generated successfully.")
    print(f"Output: {Path(args.output_docx).resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())